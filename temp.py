from openpyxl import load_workbook
from docx import Document

# Load the Excel workbook
workbook = load_workbook('data1.xlsx')
sheet = workbook.active

# Load the Word document template
document = Document('template.docx')

# Function to populate the Word document with data from Excel
def populate_document(sheet, document):
    for row in sheet.iter_rows(min_row=2, values_only=True):
        ACCT_NAME = row[6]
        DIS_AMT = row[15]
        DIS_SHDL_DATE = row[16]
        # Convert values to strings
        ACCT_NAME = str(ACCT_NAME)
        DIS_AMT = str(DIS_AMT)
        DIS_SHDL_DATE = str(DIS_SHDL_DATE)
        # Assuming the template has placeholders like "{{Name}}", "{{Age}}", "{{Email}}"
        for paragraph in document.paragraphs:
            if "{{Name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{Name}}", ACCT_NAME)
            if "{{disAmount}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disAmount}}", str(DIS_AMT))
            if "{{disDate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disDate}}", str(DIS_SHDL_DATE))
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{Name}}" in cell.text:
                        cell.text = cell.text.replace("{{Name}}", ACCT_NAME)
                    if "{{disAmount}}" in cell.text:
                       cell.text = cell.text.replace("{{disAmount}}", str(DIS_AMT))
                    if "{{disDate}}" in cell.text:
                       cell.text = cell.text.replace("{{disDate}}", str(DIS_SHDL_DATE))

# Populate the Word document with data from Excel
populate_document(sheet, document)

# Save the populated Word document
document.save('output.docx')