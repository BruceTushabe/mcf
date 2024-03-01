import pandas as pd
from openpyxl import load_workbook
from docx import Document

# Function to find data for a specific account number in the Excel sheet
def find_account_data(account_number):
    account_number = int(account_number) # Convert to string for comparison
    workbook = load_workbook('data1.xlsx', data_only=True)
    sheet = workbook.active
    data = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[2] == account_number:  # Assuming the account number is in the 6th column (index 4)
            data.append(row)
    if data:
        columns = ["Name", "Gender", "Contact", "Account number", "Age Category", "District", "Date of Loan Disbursement", "Amount disbursed", "Outstanding balance", "Amount to be recovered", "Loan cycle", "Mode of engagement", "Date of default"]
        return pd.DataFrame(data, columns=columns)
    return None  # Return None if account number is not found

# Load the Word document template
document = Document('template.docx')

# Function to populate the Word document with data for a specific account number
def populate_document_for_account(document, account_data):
    if not account_data.empty:
        ACCT_NAME = account_data['Name'].iloc[0]
        DIS_AMT = account_data['Amount disbursed'].iloc[0]
        DIS_SHDL_DATE = account_data['Date of Loan Disbursement'].iloc[0]
        Address = account_data['District'].iloc[0]
        AGE = account_data['Age Category'].iloc[0]
        Gender = account_data['Gender'].iloc[0]
        DATE_ARREARS_START = account_data['Date of default'].iloc[0]
        Application_date = account_data['Mode of engagement'].iloc[0]
        DOB = account_data['Contact'].iloc[0]
        AMOUNT_CLAIMED = account_data['Outstanding balance'].iloc[0]
        ARREARSDAYS = account_data['Amount to be recovered'].iloc[0]
        TERM = account_data['Loan cycle'].iloc[0]

        # Convert values to strings
        Address = str(Address)
        ACCT_NAME = str(ACCT_NAME)
        DIS_AMT = str(DIS_AMT)
        DIS_SHDL_DATE = str(DIS_SHDL_DATE)
        AGE = str(AGE)
        Gender = str(Gender)
        DATE_ARREARS_START = str(DATE_ARREARS_START)
        Application_date = str(Application_date)
        DOB = str(DOB)
        AMOUNT_CLAIMED = str(AMOUNT_CLAIMED)
        ARREARSDAYS = str(ARREARSDAYS)
        TERM = str(TERM)

        # Assuming the template has placeholders like "{{Name}}", "{{disAmount}}", "{{disDate}}"
        for paragraph in document.paragraphs:
            if "{{Name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{Name}}", ACCT_NAME)
            if "{{disAmount}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disAmount}}", DIS_AMT)
            if "{{disDate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disDate}}", DIS_SHDL_DATE)
            if "{{address}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{address}}", Address)
            if "{{age}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{age}}", AGE)
            if "{{gender}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{gender}}", Gender)
            if "{{dob}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{dob}}", DOB)
            if "{{appdate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{appdate}}", Application_date)
            if "{{amtclaimed}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{amtclaimed}}", AMOUNT_CLAIMED)
            if "{{datearrears}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{datearreas}}", DATE_ARREARS_START)
            if "{{arrearsdays}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{arrearsdays}}", ARREARSDAYS)
            if "{{loanterm}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{loanterm}}", TERM)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{Name}}" in cell.text:
                        cell.text = cell.text.replace("{{Name}}", ACCT_NAME)
                    if "{{disAmount}}" in cell.text:
                        cell.text = cell.text.replace("{{disAmount}}", DIS_AMT)
                    if "{{disDate}}" in cell.text:
                        cell.text = cell.text.replace("{{disDate}}", DIS_SHDL_DATE)
                    if "{{address}}" in cell.text:
                        cell.text = cell.text.replace("{{address}}", Address)
                    if "{{age}}" in cell.text:
                        cell.text = cell.text.replace("{{age}}", AGE)
                    if "{{gender}}" in cell.text:
                        cell.text = cell.text.replace("{{gender}}", Gender)
                    if "{{dob}}" in cell.text:
                        cell.text = cell.text.replace("{{dob}}", DOB)
                    if "{{appdate}}" in cell.text:
                        cell.text = cell.text.replace("{{appdate}}", Application_date)
                    if "{{amtclaimed}}" in cell.text:
                        cell.text = cell.text.replace("{{amtclaimed}}", AMOUNT_CLAIMED)
                    if "{{datearrears}}" in cell.text:
                        cell.text = cell.text.replace("{{datearreas}}", DATE_ARREARS_START)
                    if "{{arrearsdays}}" in cell.text:
                        cell.text = cell.text.replace("{{arrearsdays}}", ARREARSDAYS)
                    if "{{loanterm}}" in cell.text:
                        cell.text = cell.text.replace("{{loanterm}}", TERM)
        return True  # Data populated successfully
    return False  # Account number not found

# Input the account number
account_number = input("Enter the account number: ")

# Find data for the specified account number in the Excel sheet
account_data = find_account_data(account_number)

# Populate the Word document with data for the specified account number
if account_data is not None:
    # Populating the Word document
    if populate_document_for_account(document, account_data):
        # Save the populated Word document
        output_word_filename = f"output_word_{account_number}.docx"
        document.save(output_word_filename)
        print(f"Word document generated for account number {account_number}.")
    else:
        print(f"Failed to generate Word document for account number {account_number}.")
    
    # Creating Excel sheet
    if not account_data.empty:
        # Save the DataFrame to Excel sheet
        output_excel_filename = f"output_excel_{account_number}.xlsx"
        account_data.to_excel(output_excel_filename, index=False)
        print(f"Excel sheet generated for account number {account_number}.")
    else:
        print(f"No data found for account number {account_number} to create Excel sheet.")
else:
    print(f"Account number {account_number} not found in the Excel sheet.")
