from openpyxl import load_workbook, Workbook
from docx import Document

# Function to find data for a specific account number in the Excel sheet
def find_account_data(account_number):
    account_number = int(account_number) # Convert to string for comparison
    workbook = load_workbook('data1.xlsx', data_only=True)
    sheet = workbook.active
    for row in sheet.iter_rows(min_row=2, values_only=True):
    #    print(f"Checking row: {row}")  # Print the entire row for debugging
        if row[2] == account_number:  # Assuming the account number is in the 6th column (index 4)
            return row  # Return the row data if account number is found
    return None  # Return None if account number is not found

# Load the Word document template
document = Document('template.docx')

# Function to populate the Word document with data for a specific account number
def populate_document_for_account(document, account_data):
    if account_data:
        Address = account_data[21]
        ACCT_NAME = account_data[3]
        DIS_AMT = account_data[7]
        DIS_SHDL_DATE = account_data[8]
        AGE = account_data[16]
        Gender = account_data[20]
        DATE_ARREARS_START = account_data[10]
        Application_date = account_data[19]
        DOB = account_data[18]
        AMOUNT_CLAIMED = account_data[11]
        ARREARSDAYS = account_data[4]
        TERM = account_data[6]

        # Convert values to strings
        Address = str(Address)
        ACCT_NAME = str(ACCT_NAME)
        DIS_AMT = str(DIS_AMT)
        DIS_SHDL_DATE = str(DIS_SHDL_DATE)
        AGE = str(AGE)
        Gender = str(Gender)
        DATE_ARREARS_START = str(DATE_ARREARS_START)
        Application_date = str(Application_date)
        DOB = str(DOB)
        AMOUNT_CLAIMED = str(AMOUNT_CLAIMED)
        ARREARSDAYS = str(ARREARSDAYS)
        TERM = str(TERM)


        # Assuming the template has placeholders like "{{Name}}", "{{disAmount}}", "{{disDate}}"
        for paragraph in document.paragraphs:
            if "{{Name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{Name}}", ACCT_NAME)
            if "{{disAmount}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disAmount}}", str(DIS_AMT))
            if "{{disDate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{disDate}}", str(DIS_SHDL_DATE))
            if "{{address}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{address}}", str(Address))
            if "{{age}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{age}}", str(AGE))
            if "{{gender}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{gender}}", str(Gender))
            if "{{dob}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{dob}}", str(DOB))
            if "{{appdate}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{appdate}}", str(Application_date))
            if "{{amtclaimed}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{amtclaimed}}", str(AMOUNT_CLAIMED))
            if "{{datearrears}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{datearreas}}", str(DATE_ARREARS_START))
            if "{{arrearsdays}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{arrearsdays}}", str(ARREARSDAYS))
            if "{{loanterm}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{loanterm}}", str(TERM))
            
            
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{Name}}" in cell.text:
                        cell.text = cell.text.replace("{{Name}}", ACCT_NAME)
                    if "{{disAmount}}" in cell.text:
                        cell.text = cell.text.replace("{{disAmount}}", str(DIS_AMT))
                    if "{{disDate}}" in cell.text:
                        cell.text = cell.text.replace("{{disDate}}", str(DIS_SHDL_DATE))
                    if "{{address}}" in cell.text:
                        cell.text = cell.text.replace("{{address}}", str(Address))
                    if "{{age}}" in cell.text:
                        cell.text = cell.text.replace("{{age}}", str(AGE))
                    if "{{gender}}" in cell.text:
                        cell.text = cell.text.replace("{{gender}}", str(Gender))
                    if "{{dob}}" in cell.text:
                        cell.text = cell.text.replace("{{dob}}", str(DOB))
                    if "{{appdate}}" in cell.text:
                        cell.text = cell.text.replace("{{appdate}}", str(Application_date))
                    if "{{amtclaimed}}" in cell.text:
                        cell.text = cell.text.replace("{{amtclaimed}}", str(AMOUNT_CLAIMED))
                    if "{{datearrears}}" in cell.text:
                        cell.text = cell.text.replace("{{datearreas}}", str(DATE_ARREARS_START))
                    if "{{arrearsdays}}" in cell.text:
                        cell.text = cell.text.replace("{{arrearsdays}}", str(ARREARSDAYS))
                    if "{{loanterm}}" in cell.text:
                        cell.text = cell.text.replace("{{loanterm}}", str(TERM))
        return True  # Data populated successfully
    return False  # Account number not found

# Input the account number
account_number = input("Enter the account number: ")

# Find data for the specified account number in the Excel sheet
account_data = find_account_data(account_number)

# Populate the Word document with data for the specified account number
if populate_document_for_account(document, account_data):
    # Save the populated Word document
    output_filename = f"output_{account_number}.docx"
    document.save(output_filename)
    print(f"Word document generated for account number {account_number}.")
else:
    print(f"Account number {account_number} not found in the Excel sheet.")

# Function to pupulate excel sheet 
    

def populate_excel_from_excel(input_excel_path, output_excel_path):
    # Load data from existing Excel sheet
    input_wb = load_workbook(input_excel_path)
    input_ws = input_wb.active

    # Create a new Excel workbook and sheet
    output_wb = Workbook()
    output_ws = output_wb.active

    # Populate the new Excel sheet with data from existing Excel and Word document
    for row in input_ws.iter_rows(values_only=True):
        output_ws.append(row)

    # Save the populated Excel sheet
    output_wb.save(output_excel_path)

# Example usage
populate_excel_from_excel("data1.xlsx", "output_data.xlsx")

    

