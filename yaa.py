
import openpyxl
from openpyxl import load_workbook
from docx import Document

# Load the Excel workbook
workbook = load_workbook('data1.xlsx')
sheet = workbook.active

# Function to populate a Word document with data from a single row
def populate_document(row, template_path, output_path):
    document = Document(template_path)

    ACCT_NAME = row[6]
    DIS_AMT = row[15]
    DIS_SHDL_DATE = row[16]

    # Replace placeholders in paragraphs
    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace("{{Name}}", ACCT_NAME)
        paragraph.text = paragraph.text.replace("{{disAmount}}", str(DIS_AMT))
        paragraph.text = paragraph.text.replace("{{disDate}}", str(DIS_SHDL_DATE))

    # Replace placeholders in tables (if present)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("{{Name}}", ACCT_NAME)
                cell.text = cell.text.replace("{{disAmount}}", str(DIS_AMT))
                cell.text = cell.text.replace("{{disDate}}", str(DIS_SHDL_DATE))

    document.save(output)

# Iterate through rows and create separate documents
for row in sheet.iter_rows(min_row=2, values_only=True):
    output_filename = f"output_{row[6]}.docx"  # Use a unique filename based on name or other relevant field
    populate_document(row, "template.docx", output_filename)